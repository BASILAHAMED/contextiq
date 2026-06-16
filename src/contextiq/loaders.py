from __future__ import annotations

import csv
import json
from html.parser import HTMLParser
from pathlib import Path

from .models import Document, Section
from .utils import normalize_text, stable_id


class PlainTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data.strip())

    def get_text(self) -> str:
        return "\n".join(self.parts)


def extract_sections(text: str) -> list[Section]:
    sections: list[Section] = []
    lines = text.splitlines()
    cursor = 0
    heading_positions: list[tuple[str, int, int]] = []
    for line in lines:
        stripped = line.strip()
        line_len = len(line) + 1
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip() or "Section"
            heading_positions.append((title, level, cursor))
        cursor += line_len

    for index, (title, level, start) in enumerate(heading_positions):
        end = heading_positions[index + 1][2] if index + 1 < len(heading_positions) else len(text)
        sections.append(Section(title=title, level=level, start_char=start, end_char=end))
    return sections


def load_document(path: Path, root: Path) -> Document:
    suffix = path.suffix.lower()
    rel_path = path.relative_to(root).as_posix()
    warnings: list[str] = []

    if suffix in {".txt", ".md", ".rst"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        source_type = suffix[1:]
    elif suffix in {".json", ".jsonl"}:
        text = _load_json_like(path, suffix)
        source_type = suffix[1:]
    elif suffix in {".csv", ".tsv"}:
        text = _load_delimited(path, "\t" if suffix == ".tsv" else ",")
        source_type = suffix[1:]
    elif suffix in {".html", ".htm"}:
        text = _load_html(path)
        source_type = "html"
    elif suffix == ".pdf":
        text, warning = _load_pdf(path)
        source_type = "pdf"
        if warning:
            warnings.append(warning)
    elif suffix == ".docx":
        text, warning = _load_docx(path)
        source_type = "docx"
        if warning:
            warnings.append(warning)
    else:
        raise ValueError(f"Unsupported file type: {path}")

    normalized = normalize_text(text)
    return Document(
        doc_id=stable_id(rel_path),
        source_path=rel_path,
        source_type=source_type,
        text=normalized,
        metadata={"extension": suffix},
        sections=extract_sections(normalized),
        warnings=warnings,
    )


def _load_json_like(path: Path, suffix: str) -> str:
    if suffix == ".json":
        data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        return json.dumps(data, indent=2, ensure_ascii=False)

    rows = []
    with path.open("r", encoding="utf-8", errors="ignore") as handle:
        for line in handle:
            stripped = line.strip()
            if stripped:
                rows.append(json.loads(stripped))
    return "\n\n".join(json.dumps(row, indent=2, ensure_ascii=False) for row in rows)


def _load_delimited(path: Path, delimiter: str) -> str:
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        lines = [" | ".join(cell.strip() for cell in row) for row in reader]
    return "\n".join(lines)


def _load_html(path: Path) -> str:
    parser = PlainTextExtractor()
    parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
    return parser.get_text()


def _load_pdf(path: Path) -> tuple[str, str | None]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return (
            "",
            "Skipped PDF content because 'pypdf' is not installed. Install with: pip install contextiq[docs]",
        )

    reader = PdfReader(str(path))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(parts), None


def _load_docx(path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document as WordDocument
    except ImportError:
        return (
            "",
            "Skipped DOCX content because 'python-docx' is not installed. Install with: pip install contextiq[docs]",
        )

    doc = WordDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs), None
